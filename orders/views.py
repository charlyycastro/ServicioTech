import os
import uuid
import base64
import unicodedata
import ast # Vital para leer listas de IDs
import re
import markdown
# --- TERCEROS ---
import weasyprint
from google import genai
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# --- DJANGO ---
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.views.decorators.http import require_POST
from django.views.decorators.cache import never_cache
from django.urls import reverse_lazy
from django.contrib import messages
from django.contrib.auth import logout
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.db.models import Q
from django.conf import settings
from django.core.paginator import Paginator
from django.contrib.auth.models import User
from django.contrib.auth.hashers import make_password
from django.core.mail import EmailMessage
from django.core.files.base import ContentFile
from django.utils.dateparse import parse_date
from django.utils import timezone
from django.db import IntegrityError, transaction
from django.views.generic import ListView, CreateView, UpdateView, DeleteView, DetailView
from django.contrib.auth.mixins import LoginRequiredMixin


# --- MODELOS Y FORMULARIOS ---
from .models import ServiceOrder, EngineerProfile, ServiceEvidence, TechnicalMemory 
from .forms import (
    ServiceOrderForm, EquipmentFormSet, ServiceMaterialFormSet,
    ShelterEquipmentFormSet, ServiceEvidenceFormSet,
    CustomUserCreationForm, UserEditForm
)
# ================================================================
# FUNCIONES AUXILIARES
# ================================================================

def es_ingeniero_o_admin(user):
    """Verifica si el usuario es Staff (Ingeniero) o Superusuario"""
    return user.is_authenticated and (user.is_staff or user.is_superuser)

def es_superusuario(user):
    """Verifica si el usuario es Superusuario (Admin total)"""
    return user.is_authenticated and user.is_superuser

def obtener_firma_ingeniero(nombre_ingeniero, request):
    """Busca el perfil del ingeniero por nombre completo y retorna la URL absoluta de su firma."""
    if not nombre_ingeniero:
        return None
    
    # Buscamos quién tiene ese nombre completo
    for user in User.objects.all():
        full_name = user.get_full_name()
        # Comparación flexible (nombre completo o username si no tiene nombre)
        if full_name == nombre_ingeniero or user.username == nombre_ingeniero:
            if hasattr(user, 'profile') and user.profile.firma:
                # build_absolute_uri es vital para que WeasyPrint cargue la imagen en el PDF
                return request.build_absolute_uri(user.profile.firma.url)
    return None

def guardar_firma(user, data_url):
    """Decodifica el string base64 del canvas y lo guarda como imagen en el perfil."""
    try:
        if ';base64,' in data_url:
            format, imgstr = data_url.split(';base64,') 
            ext = format.split('/')[-1] 
            data = ContentFile(base64.b64decode(imgstr), name=f'firma_{user.username}_{uuid.uuid4()}.{ext}')
            
            profile, created = EngineerProfile.objects.get_or_create(user=user)
            profile.firma = data
            profile.save()
    except Exception as e:
        print(f"Error guardando firma de usuario: {e}")


# ================================================================
# DASHBOARD (CEREBRO DEL SISTEMA)
# ================================================================

@login_required
def dashboard_view(request):
    user = request.user
    nombre_usuario = user.get_full_name() or user.username

    # 1. TOTAL HISTÓRICO (Global de la empresa)
    # Cuenta TODO lo que existe en la base de datos, sin importar de quién sea.
    total = ServiceOrder.objects.count()

    # 2. PENDIENTES (Global de la empresa)
    # Cuenta todas las que están en pausa/borrador.
    pendientes = ServiceOrder.objects.filter(estatus='borrador').count()

    # 3. FINALIZADAS (Global de la empresa)
    # Cuenta todas las que ya se cerraron.
    finalizadas = ServiceOrder.objects.filter(estatus='finalizado').count()

    # 4. MIS ASIGNACIONES (Personal)
    # Aquí sí filtramos por TU nombre y que NO estén terminadas.
    mis_asignadas = ServiceOrder.objects.filter(
        ingeniero_nombre__icontains=nombre_usuario
    ).exclude(estatus='finalizado').count()

    # 5. TABLA RECIENTES (Global)
    # Muestra las últimas 5 que han entrado al sistema.
    recientes = ServiceOrder.objects.all().order_by('-creado')[:5]

    context = {
        'total': total,
        'pendientes': pendientes,
        'finalizadas': finalizadas,
        'mis_asignadas': mis_asignadas, # Esta será 0 si nadie te ha asignado nada a ti "Carlos"
        'recientes': recientes,
    }
    return render(request, 'orders/dashboard.html', context)


# ================================================================
# VISTAS DE ÓRDENES (CRUD)
# ================================================================

@login_required
def order_list(request):
    orders = ServiceOrder.objects.all().order_by('-creado')

    # Filtros
    query = request.GET.get('q', '').strip()
    filtro_empresa = request.GET.get('empresa')
    filtro_estatus = request.GET.get('estatus')
    filtro_ingeniero = request.GET.get('ingeniero')
    fecha_inicio = request.GET.get('fecha_inicio')
    fecha_fin = request.GET.get('fecha_fin')

    if query:
        orders = orders.filter(
            Q(folio__icontains=query) |
            Q(cliente_nombre__icontains=query) |  
            Q(cliente_contacto__icontains=query) | 
            Q(titulo__icontains=query)
        )

    if filtro_empresa:
        orders = orders.filter(cliente_nombre=filtro_empresa)
    if filtro_estatus:
        orders = orders.filter(estatus=filtro_estatus)
    if filtro_ingeniero:
        orders = orders.filter(ingeniero_nombre=filtro_ingeniero)
    if fecha_inicio:
        orders = orders.filter(creado__date__gte=fecha_inicio) 
    if fecha_fin:
        orders = orders.filter(creado__date__lte=fecha_fin)

    # Listas para selects de filtro
    empresas = ServiceOrder.objects.exclude(cliente_nombre__isnull=True).exclude(cliente_nombre__exact='').values_list('cliente_nombre', flat=True).distinct().order_by('cliente_nombre')
    ingenieros_list = ServiceOrder.objects.exclude(ingeniero_nombre__isnull=True).exclude(ingeniero_nombre__exact='').values_list('ingeniero_nombre', flat=True).distinct().order_by('ingeniero_nombre')

    # Paginación
    paginator = Paginator(orders, 15)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    ctx = {
        'page_obj': page_obj, 
        'query': query,
        'empresas': empresas,
        'ingenieros_list': ingenieros_list,
        'filtro_empresa': filtro_empresa,
        'filtro_estatus': filtro_estatus,
        'filtro_ingeniero': filtro_ingeniero,
        'fecha_inicio': fecha_inicio,
        'fecha_fin': fecha_fin,       
        'STATUS_CHOICES': ServiceOrder.STATUS_CHOICES 
    }
    return render(request, 'orders/order_list.html', ctx)

@login_required
def order_detail(request, pk):
    order = get_object_or_404(ServiceOrder, pk=pk)
    
    # 1. Firma Ingeniero (Lógica existente)
    firma_ingeniero_url = obtener_firma_ingeniero(order.ingeniero_nombre, request)

    # 2. NUEVO: Firma Visor (Ventas)
    firma_visor_url = None
    # Verificamos si hay un usuario seleccionado en el campo 'visor' y si tiene perfil con firma
    if order.visor and hasattr(order.visor, 'profile') and order.visor.profile.firma:
        # build_absolute_uri es CRUCIAL para que WeasyPrint vea la imagen
        firma_visor_url = request.build_absolute_uri(order.visor.profile.firma.url)

    return render(request, 'orders/order_detail.html', {
        'object': order,
        'firma_ingeniero_url': firma_ingeniero_url,
        'firma_visor_url': firma_visor_url, # <--- Agregamos esto al contexto
    })

# ===== CREAR ORDEN =====
@login_required
@user_passes_test(es_ingeniero_o_admin)
def order_create(request):
    if request.method == "POST":
        form = ServiceOrderForm(request.POST, request.FILES)
        equipos_fs = EquipmentFormSet(request.POST, request.FILES, prefix="equipos")
        materiales_fs = ServiceMaterialFormSet(request.POST, request.FILES, prefix="materiales")
        resguardos_fs = ShelterEquipmentFormSet(request.POST, request.FILES, prefix="resguardos")
        evidencias_fs = ServiceEvidenceFormSet(request.POST, request.FILES, prefix="evidencias")

        if form.is_valid() and equipos_fs.is_valid() and materiales_fs.is_valid() and resguardos_fs.is_valid() and evidencias_fs.is_valid():
            order = form.save(commit=False)
            accion = request.POST.get('accion', 'borrador')

            if accion == 'finalizar':
                errores = []
                if not order.cliente_contacto: errores.append("Persona de contacto")
                if not order.cliente_email: errores.append("Correo del cliente")
                if not order.ingeniero_nombre: errores.append("Ingeniero asignado")
                if not order.tipos_servicio: errores.append("Tipo de servicio")
                
                if errores:
                    messages.error(request, f"Para finalizar, faltan: {', '.join(errores)}")
                    ctx = {
                        "form": form, "equipos_fs": equipos_fs, "materiales_fs": materiales_fs,
                        "resguardos_fs": resguardos_fs, "evidencias_fs": evidencias_fs,
                        "titulo": "Nueva Orden de Servicio"
                    }
                    return render(request, "orders/order_form.html", ctx)

                order.estatus = 'finalizado'
                mensaje_exito = f"Orden {order.folio} FINALIZADA correctamente."
            else:
                order.estatus = 'borrador'
                mensaje_exito = f"Borrador guardado."

            # Firma Cliente
            firma_b64 = (request.POST.get("firma_b64") or "").strip()
            if firma_b64.startswith("data:image"):
                try:
                    header, data = firma_b64.split(",", 1)
                    ext = "png"
                    if "jpeg" in header.lower(): ext = "jpg"
                    file_data = ContentFile(base64.b64decode(data), name=f"firma_cliente_{uuid.uuid4().hex}.{ext}")
                    order.firma = file_data
                except Exception as e:
                    print(f"Error firma cliente: {e}")

            order.save()
            
            # --- CARGA MASIVA DE FOTOS (NUEVO) ---
            archivos_masivos = request.FILES.getlist('imagenes_masivas')
            for f in archivos_masivos:
                ServiceEvidence.objects.create(order=order, archivo=f)
            # -------------------------------------

            equipos_fs.instance = order; equipos_fs.save()
            materiales_fs.instance = order; materiales_fs.save()
            resguardos_fs.instance = order; resguardos_fs.save()
            evidencias_fs.instance = order; evidencias_fs.save()

            messages.success(request, mensaje_exito)
            return redirect("orders:detail", pk=order.pk)
        else:
            messages.error(request, "Hay errores en el formulario.")
    else:
        initial_data = {}
        if request.user.is_staff:
             initial_data['ingeniero_nombre'] = request.user.get_full_name() or request.user.username

        form = ServiceOrderForm(initial=initial_data)
        equipos_fs = EquipmentFormSet(prefix="equipos")
        materiales_fs = ServiceMaterialFormSet(prefix="materiales")
        resguardos_fs = ShelterEquipmentFormSet(prefix="resguardos")
        evidencias_fs = ServiceEvidenceFormSet(prefix="evidencias")

    ctx = {
        "form": form, "equipos_fs": equipos_fs, "materiales_fs": materiales_fs,
        "resguardos_fs": resguardos_fs, "evidencias_fs": evidencias_fs,
        "titulo": "Nueva Orden de Servicio"
    }
    return render(request, "orders/order_form.html", ctx)


# ===== EDITAR ORDEN =====
@login_required
@user_passes_test(es_ingeniero_o_admin)
def order_update(request, pk):
    order = get_object_or_404(ServiceOrder, pk=pk)
    # --- BLOQUEO DE SEGURIDAD INTELIGENTE ---
    # Si el usuario NO es Superusuario Y la orden YA está finalizada...
    if not request.user.is_superuser and order.estatus == 'finalizado':
        messages.error(request, "⚠️ No puedes editar una orden finalizada. Solicita correcciones al Administrador.")
        return redirect('orders:detail', pk=pk) # Lo regresamos al detalle
    # ----------------------------------------
    if request.method == "POST":
        form = ServiceOrderForm(request.POST, request.FILES, instance=order)
        equipos_fs = EquipmentFormSet(request.POST, request.FILES, instance=order, prefix="equipos")
        materiales_fs = ServiceMaterialFormSet(request.POST, request.FILES, instance=order, prefix="materiales")
        resguardos_fs = ShelterEquipmentFormSet(request.POST, request.FILES, instance=order, prefix="resguardos")
        evidencias_fs = ServiceEvidenceFormSet(request.POST, request.FILES, instance=order, prefix="evidencias")

        if form.is_valid() and equipos_fs.is_valid() and materiales_fs.is_valid() and resguardos_fs.is_valid() and evidencias_fs.is_valid():
            order = form.save(commit=False)
            accion = request.POST.get('accion', 'borrador')

            if accion == 'finalizar':
                errores = []
                if not order.cliente_contacto: errores.append("Persona de contacto")
                if not order.cliente_email: errores.append("Correo del cliente")
                if not order.ingeniero_nombre: errores.append("Ingeniero asignado")
                if not order.tipos_servicio: errores.append("Tipo de servicio")
                
                if errores:
                    messages.error(request, f"Para finalizar, faltan: {', '.join(errores)}")
                    ctx = {
                        "form": form, "equipos_fs": equipos_fs, "materiales_fs": materiales_fs,
                        "resguardos_fs": resguardos_fs, "evidencias_fs": evidencias_fs,
                        "titulo": f"Editar Orden {order.folio}"
                    }
                    return render(request, "orders/order_form.html", ctx)

                order.estatus = 'finalizado'
                mensaje = f"Orden {order.folio} FINALIZADA."
            else:
                order.estatus = 'borrador'
                mensaje = f"Cambios guardados."

            firma_b64 = (request.POST.get("firma_b64") or "").strip()
            if firma_b64.startswith("data:image"):
                try:
                    header, data = firma_b64.split(",", 1)
                    ext = "png" if "png" in header.lower() else "jpg"
                    file_data = ContentFile(base64.b64decode(data), name=f"firma_cliente_{uuid.uuid4().hex}.{ext}")
                    order.firma = file_data
                except Exception:
                    pass

            order.save()
            
            # --- CARGA MASIVA DE FOTOS (NUEVO) ---
            archivos_masivos = request.FILES.getlist('imagenes_masivas')
            for f in archivos_masivos:
                ServiceEvidence.objects.create(order=order, archivo=f)
            # -------------------------------------

            equipos_fs.save()
            materiales_fs.save()
            resguardos_fs.save()
            evidencias_fs.save()

            messages.success(request, mensaje)
            return redirect("orders:detail", pk=order.pk)
        else:
            messages.error(request, "Hay errores en el formulario.")
    else:
        form = ServiceOrderForm(instance=order)
        equipos_fs = EquipmentFormSet(instance=order, prefix="equipos")
        materiales_fs = ServiceMaterialFormSet(instance=order, prefix="materiales")
        resguardos_fs = ShelterEquipmentFormSet(instance=order, prefix="resguardos")
        evidencias_fs = ServiceEvidenceFormSet(instance=order, prefix="evidencias")

    ctx = {
        "form": form,
        "equipos_fs": equipos_fs, "materiales_fs": materiales_fs,
        "resguardos_fs": resguardos_fs, "evidencias_fs": evidencias_fs,
        "titulo": f"Editar Orden {order.folio}"
    }
    return render(request, "orders/order_form.html", ctx)


# ===== ELIMINAR ORDENES =====
@require_POST
@login_required
@user_passes_test(es_superusuario)
def bulk_delete(request):
    ids = request.POST.getlist("ids") or request.POST.getlist("selected")
    if not ids:
        messages.warning(request, "No seleccionaste ninguna orden.")
        return redirect("orders:list")

    qs = ServiceOrder.objects.filter(id__in=ids)
    count = qs.count()
    qs.delete()

    messages.success(request, f"Se eliminaron {count} órdenes correctamente.")
    return redirect("orders:list")


# ===== ENVIAR CORREO =====
@login_required
@user_passes_test(es_ingeniero_o_admin)
def email_order(request, pk):
    order = get_object_or_404(ServiceOrder, pk=pk)

    if not order.cliente_email:
        messages.warning(request, "Esta orden no tiene un correo de cliente registrado.")
        return redirect('orders:detail', pk=pk)

    firma_ingeniero_url = obtener_firma_ingeniero(order.ingeniero_nombre, request)
    if not firma_ingeniero_url:
        messages.error(request, f"No se puede enviar. El ingeniero '{order.ingeniero_nombre}' no tiene firma precargada.")
        return redirect('orders:detail', pk=pk)

    survey_link = "https://www.cognitoforms.com/INOVATECH1/EncuestaDeServicio"

    try:
        import weasyprint
        import smtplib
        import ssl
        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText
        from email.mime.application import MIMEApplication

        html_string = render_to_string('orders/order_detail.html', {
            'object': order,
            'print_mode': True,
            'firma_ingeniero_url': firma_ingeniero_url 
        }, request=request)

        pdf_bytes = weasyprint.HTML(
            string=html_string,
            base_url=request.build_absolute_uri()
        ).write_pdf()

        # Configuración SMTP
        smtp_server = settings.EMAIL_HOST
        smtp_port = settings.EMAIL_PORT
        smtp_user = settings.EMAIL_HOST_USER
        smtp_password = settings.EMAIL_HOST_PASSWORD

        msg = MIMEMultipart()
        msg['From'] = settings.DEFAULT_FROM_EMAIL
        msg['To'] = order.cliente_email
        msg['Subject'] = f"Reporte de Servicio {order.folio} - ServicioTech"

        body = (
            f"Hola {order.cliente_contacto or 'Cliente'},\n\n"
            f"Adjunto encontrarás el reporte de servicio técnico realizado el {order.fecha_servicio}.\n\n"
            f"Folio: {order.folio}\n"
            f"Título: {order.titulo}\n\n"
            "=============================================\n"
            "¡AYÚDANOS A MEJORAR!\n"
            "=============================================\n"
            "Por favor, tómate un minuto para calificar nuestro servicio:\n"
            f"{survey_link}\n\n" 
            "Atentamente,\n"
            "El equipo de ServicioTech."
        )
        msg.attach(MIMEText(body, 'plain'))
        
        attachment = MIMEApplication(pdf_bytes, _subtype="pdf")
        attachment.add_header('Content-Disposition', 'attachment', filename=f"{order.folio}.pdf")
        msg.attach(attachment)

        context = ssl.create_default_context()
        context.check_hostname = False
        context.verify_mode = ssl.CERT_NONE

        server = smtplib.SMTP(smtp_server, smtp_port, timeout=30)
        server.starttls(context=context)
        server.login(smtp_user, smtp_password)
        server.send_message(msg)
        server.quit()

        order.email_enviado = True
        order.save(update_fields=['email_enviado'])
        messages.success(request, f"Correo enviado correctamente a {order.cliente_email}")

    except Exception as e:
        messages.error(request, f"Error de conexión: {e}")

    return redirect('orders:detail', pk=pk)


# ================================================================
# GESTIÓN DE USUARIOS
# ================================================================

@login_required
@user_passes_test(es_superusuario)
def user_list_view(request):
    users = User.objects.all().order_by('username')
    return render(request, 'orders/user_list.html', {'users': users})

@login_required
@user_passes_test(es_superusuario)
def create_user_view(request):
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.password = make_password(form.cleaned_data['password'])
            
            role = form.cleaned_data['role']
            user.is_staff = (role in ['ingeniero', 'superuser'])
            user.is_superuser = (role == 'superuser')
            user.save()

            firma_b64 = form.cleaned_data.get('firma_b64')
            if firma_b64 and role in ['ingeniero', 'visor']:
                    guardar_firma(user, firma_b64)

            messages.success(request, f"Usuario {user.username} creado.")
            return redirect('orders:user_list')
    else:
        form = CustomUserCreationForm()
    
    return render(request, 'orders/user_form.html', {'form': form, 'titulo': 'Alta de Nuevo Usuario'})

@login_required
@user_passes_test(es_superusuario)
def edit_user_view(request, pk):
    user_to_edit = get_object_or_404(User, pk=pk)
    
    rol_inicial = 'visor'
    if user_to_edit.is_superuser: rol_inicial = 'superuser'
    elif user_to_edit.is_staff: rol_inicial = 'ingeniero'

    if request.method == 'POST':
        form = UserEditForm(request.POST, instance=user_to_edit)
        if form.is_valid():
            user = form.save(commit=False)
            if form.cleaned_data['password']:
                user.password = make_password(form.cleaned_data['password'])
            
            role = form.cleaned_data['role']
            user.is_staff = (role in ['ingeniero', 'superuser'])
            user.is_superuser = (role == 'superuser')
            user.save()

            firma_b64 = form.cleaned_data.get('firma_b64')
            if firma_b64 and role in ['ingeniero', 'visor']:
                          guardar_firma(user, firma_b64)

            messages.success(request, f"Usuario {user.username} actualizado.")
            return redirect('orders:user_list')
    else:
        form = UserEditForm(instance=user_to_edit, initial={'role': rol_inicial})

    return render(request, 'orders/user_form.html', {'form': form, 'titulo': f'Editar a {user_to_edit.username}'})

@login_required
@user_passes_test(es_superusuario)
def delete_user_view(request, pk):
    usuario_a_borrar = get_object_or_404(User, pk=pk)
    
    if usuario_a_borrar.is_superuser:
        messages.error(request, "No puedes eliminar al Superadministrador.")
        return redirect('orders:user_list')

    try:
        nombre = usuario_a_borrar.username
        usuario_a_borrar.delete()
        messages.success(request, f"El usuario '{nombre}' fue eliminado.")
        
    except IntegrityError:
        messages.error(request, f"No se puede eliminar a '{usuario_a_borrar.username}' porque tiene historial. Desactívalo.")
    except Exception as e:
        messages.error(request, f"Error: {e}")

    return redirect('orders:user_list')


# ================================================================
# LOGIN / LOGOUT
# ================================================================

def login_view(request):
    return redirect('account_login')

@require_POST
@never_cache 
def logout_view(request):
    logout(request)
    messages.info(request, "Has cerrado sesión.")
    return redirect("account_login")



# ================================================================
# AUXILIARES PARA EL WORD (DISEÑO CORPORATIVO)
# ================================================================

def add_page_number(run):
    """Inserta numeración dinámica 'Página X'."""
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def add_table_of_contents(paragraph):
    """Inserta el código de campo para el Índice Automático (TOC)."""
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def add_formatted_text(paragraph, text):
    """Parsea negritas (**) para aplicarlas en el Word."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part.replace('**', ''))
            run.bold = True
        else:
            paragraph.add_run(part)

# ================================================================
# VISTA 1: SELECCIÓN (RESTAURADA PARA EVITAR ATTRIBUTEERROR)
# ================================================================

@login_required
def memory_selection_view(request):
    """Muestra órdenes finalizadas. Variable 'orders' para el HTML."""
    orders_query = ServiceOrder.objects.filter(estatus='finalizado').order_by('-fecha_servicio')
    
    cliente_f = request.GET.get('cliente', '').strip()
    if cliente_f:
        orders_query = orders_query.filter(cliente_nombre__icontains=cliente_f)
    
    clientes_list = ServiceOrder.objects.values_list('cliente_nombre', flat=True).distinct().order_by('cliente_nombre')

    # 'orders' coincide con {% for orden in orders %} de tu plantilla
    return render(request, 'orders/memory_selection.html', {
        'orders': orders_query, 
        'clientes_list': clientes_list,
        'titulo': 'Selección de Órdenes'
    })

# ================================================================
# VISTA 2: PREVIEW IA (MODELO FLASH-LATEST)
# ================================================================

# orders/views.py

@login_required
@transaction.atomic
def memory_preview_view(request):
    """VISTA 2: Previsualización. Asegura que el Historial sea visible."""
    if request.method != "POST":
        return redirect('orders:memory_select')

    selected_ids = request.POST.getlist('selected_orders')
    ordenes = ServiceOrder.objects.filter(id__in=selected_ids).order_by('fecha_servicio')
    
    if not ordenes.exists():
        messages.error(request, "Selecciona al menos una orden.")
        return redirect('orders:memory_select')

    # Limpiamos el nombre del cliente para evitar errores de búsqueda
    cliente_raw = ordenes.first().cliente_nombre
    cliente_busqueda = cliente_raw.strip()

    # Generación con IA (gemini-flash-latest)
    contexto = "".join([f"\n[FOLIO {o.folio}] Actividades: {o.actividades}." for o in ordenes])
    prompt = f"Actúa como Consultor Senior. Redacta una MEMORIA TÉCNICA para {cliente_busqueda}. Usa Markdown. Datos: {contexto}"
    
    try:
        client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))
        response = client.models.generate_content(model="gemini-flash-latest", contents=prompt)
        texto_ia = response.text if response.text else "Redacte manualmente."
    except Exception:
        texto_ia = "Error de conexión con Gemini."

    texto_html = markdown.markdown(texto_ia)

    # Persistencia v1
    memoria_db = TechnicalMemory.objects.create(
        cliente_nombre=cliente_busqueda,
        contenido_v1_ia=texto_html,
        creado_por=request.user
    )
    memoria_db.orders.set(ordenes)

    # BUSQUEDA DEL HISTORIAL (Solución para image_fa1e85.png)
    # Buscamos registros previos ignorando mayúsculas/minúsculas
    historial_query = TechnicalMemory.objects.filter(
        cliente_nombre__iexact=cliente_busqueda
    ).order_by('-creado')[:5]

    return render(request, 'orders/memory_preview.html', {
        'texto_generado': texto_html,
        'memoria_id': memoria_db.id,
        'selected_ids': selected_ids,
        'cliente': cliente_busqueda,
        'historial': historial_query, # <--- Esta variable llena el 'Historial de Control'
    })

@login_required
@require_POST
def memory_download_view(request):
    memoria_id = request.POST.get('memoria_id')
    texto_html = request.POST.get('texto_final', '')
    selected_ids_str = request.POST.get('selected_ids', '[]')
    
    # 1. Variables de Respaldo y Persistencia
    fecha_ejecucion = "No definida"
    tecnico = "No asignado"
    
    texto_pre_limpio = texto_html.replace('</p>', '\n').replace('<br>', '\n').replace('</h2>', '\n')
    texto_limpio = re.sub('<[^<]+?>', '', texto_pre_limpio).strip()

    memoria = get_object_or_404(TechnicalMemory, id=memoria_id)
    memoria.contenido_v2_user = texto_limpio
    memoria.save()

    try:
        selected_ids = ast.literal_eval(selected_ids_str)
        ordenes = ServiceOrder.objects.filter(id__in=selected_ids).order_by('fecha_servicio')
        orden_principal = ordenes.first()
        
        if orden_principal:
            if orden_principal.fecha_servicio:
                fecha_ejecucion = orden_principal.fecha_servicio.strftime("%d/%m/%Y")
            if orden_principal.ingeniero_nombre:
                tecnico = orden_principal.ingeniero_nombre
    except: 
        return redirect('orders:memory_select')

    document = Document()

    # --- ENCABEZADO Y PIE (Logo Inova y Páginas) ---
    section = document.sections[0]
    header_para = section.header.paragraphs[0]
    header_para.alignment = 2 # Derecha
    logo_path = os.path.join(settings.BASE_DIR, 'static', 'orders', 'inovatech-logo.png')
    if os.path.exists(logo_path):
        header_para.add_run().add_picture(logo_path, width=Inches(1.0))

    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = 1 # Centro
    footer_para.add_run("Página ")
    add_page_number(footer_para.add_run())

    # --- PORTADA (RESTAURADA SEGÚN TU DISEÑO) ---
    meses_es = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
    mes_es = meses_es[timezone.now().month - 1]
    
    document.add_paragraph("\n" * 4)
    p_titulo = document.add_paragraph()
    p_titulo.alignment = 1
    run_t = p_titulo.add_run("MEMORIA TÉCNICA")
    run_t.bold = True; run_t.font.size = Pt(28); run_t.font.color.rgb = RGBColor(31, 78, 120)

    document.add_paragraph("\n" * 2)
    p_cliente = document.add_paragraph()
    p_cliente.alignment = 1
    run_c = p_cliente.add_run(memoria.cliente_nombre.upper())
    run_c.bold = True; run_c.font.size = Pt(22)
    
    p_sub = document.add_paragraph("Servicio de Ingeniería y Soporte Técnico")
    p_sub.alignment = 1; p_sub.runs[0].font.size = Pt(14)

    document.add_paragraph("\n" * 7)
    p_empresa = document.add_paragraph("IT SOLUCIONES DE INNOVACION TECNOLOGICA AVANZA SA DE CV")
    p_empresa.alignment = 1; p_empresa.runs[0].bold = True
    
    p_loc = document.add_paragraph(f"Monterrey, N.L. México | {mes_es} {timezone.now().year}")
    p_loc.alignment = 1
    document.add_page_break()

    # --- TABLAS DE CONTROL ---
    document.add_heading('CONTROL DE DOCUMENTACIÓN', level=1)
    table_info = document.add_table(rows=3, cols=2); table_info.style = 'Table Grid'
    for i, (l, v) in enumerate([("Autor:", request.user.get_full_name() or request.user.username), ("Área:", "Ingeniería de IT"), ("Localidad:", "Monterrey, Nuevo León")]):
        table_info.cell(i, 0).text = l; table_info.cell(i, 1).text = v

    document.add_paragraph("\n")
    document.add_heading('HISTORIAL DE VERSIONES', level=1)
    table_hist = document.add_table(rows=2, cols=5); table_hist.style = 'Table Grid'
    for i, h in enumerate(["Versión", "Fecha", "Nombre", "Estado", "Comentario"]): table_hist.cell(0, i).text = h
    row = table_hist.rows[1].cells
    row[0].text = "2.0" if memoria.contenido_v2_user else "1.0"
    row[1].text = timezone.now().strftime("%d/%m/%Y"); row[2].text = tecnico
    row[3].text = "Finalizado"; row[4].text = "Revisión y firma"
    document.add_page_break()

    # --- TABLA DE CONTENIDOS (REEMPLAZA AL ÍNDICE) ---
    document.add_heading('TABLA DE CONTENIDOS', level=1)
    add_table_of_contents(document.add_paragraph())
    document.add_page_break()

    # --- RESUMEN DEL SERVICIO ---
    document.add_heading('RESUMEN DEL SERVICIO', level=1)
    table_sum = document.add_table(rows=5, cols=2); table_sum.style = 'Table Grid'
    resumen_data = [
        ("Folio de Referencia:", orden_principal.folio if orden_principal else "N/A"),
        ("Cliente:", memoria.cliente_nombre),
        ("Ubicación del Servicio:", orden_principal.ubicacion if (orden_principal and orden_principal.ubicacion) else "No especificada"),
        ("Fecha de Ejecución:", fecha_ejecucion),
        ("Técnico Responsable:", tecnico)
    ]
    for i, (l, v) in enumerate(resumen_data):
        table_sum.cell(i, 0).text = l; table_sum.cell(i, 1).text = str(v)
    document.add_paragraph("\n")

    # --- CUERPO TÉCNICO (Detección refinada de Títulos) ---
    for line in texto_limpio.split('\n'):
        line = line.strip()
        if not line: continue
        
        # Regla: Solo números romanos I al X al inicio de línea son Heading 2
        # Quitamos la regla de 'isupper' para que no tome textos normales como títulos
        if re.match(r'^(I|II|III|IV|V|VI|VII|VIII|IX|X)\.\s', line):
            document.add_heading(line, level=2)
        # Regla: A., B., C. o 1., 2. son Heading 3
        elif re.match(r'^([A-Z]|\d+)\.\s', line):
            document.add_heading(line, level=3)
        else:
            p = document.add_paragraph(); p.alignment = 3
            add_formatted_text(p, line)

    # --- ANEXO FOTOGRÁFICO (ANTES DE ELABORO) ---
    evidencias = []
    for o in ordenes: evidencias.extend(o.evidencias.all())

    if evidencias:
        document.add_page_break()
        document.add_heading('ANEXO FOTOGRÁFICO', level=1)
        table_pics = document.add_table(rows=0, cols=2)
        for i in range(0, len(evidencias), 2):
            row_cells = table_pics.add_row().cells
            for j in range(2):
                if i + j < len(evidencias):
                    ev = evidencias[i+j]
                    if ev.archivo and os.path.exists(ev.archivo.path):
                        p = row_cells[j].paragraphs[0]; p.alignment = 1
                        try:
                            run = p.add_run(); run.add_picture(ev.archivo.path, width=Inches(2.5))
                            cap = row_cells[j].add_paragraph(f"Fig {i+j+1}. {ev.comentario or 'Evidencia'}")
                            cap.alignment = 1; cap.runs[0].font.size = Pt(9)
                        except: pass

    # --- FIRMA (ELABORO) ---
    document.add_page_break()
    document.add_paragraph("\n" * 5)
    f = document.add_paragraph("ELABORO:"); f.alignment = 1
    f.add_run("\n" * 3 + f"{tecnico.upper()}\n").bold = True
    f.add_run("Gerencia Técnica")

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="Memoria_{memoria.cliente_nombre.replace(" ", "_")}.docx"'
    document.save(response)
    return response

    # ================================================================
# RESTAURACIÓN: FUNCIONES PARA WORD INDIVIDUAL (REPORTES CLÁSICOS)
# ================================================================

# --- Helpers para estilos de Word ---
def set_cell_color(cell, color):
    tc = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tc.append(shd)

def make_header_blue(cell, text):
    cell.text = ""; p = cell.paragraphs[0]; r = p.add_run(text)
    r.bold = True; r.font.color.rgb = RGBColor(255,255,255); set_cell_color(cell, '1F4E78')

def make_label_gray(cell, text):
    cell.text = ""; p = cell.paragraphs[0]; r = p.add_run(text)
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0,0,0); set_cell_color(cell, 'F2F2F2')

def set_value_text(cell, text):
    cell.text = str(text) if text else "-"; cell.paragraphs[0].runs[0].font.size = Pt(9)

def insert_signature(cell, title, img_field, name):
    cell.text = ""; p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if img_field:
        try:
            if hasattr(img_field, 'path') and os.path.exists(img_field.path):
                p.add_run().add_picture(img_field.path, height=Inches(0.6))
            elif isinstance(img_field, str) and os.path.exists(img_field):
                p.add_run().add_picture(img_field, height=Inches(0.6))
        except: pass
    else: p.add_run("\n\n\n")
    p.add_run("\n_______________________\n")
    if name: p.add_run(str(name)+"\n").bold=True
    p.add_run(title).font.size = Pt(7)

@login_required
def download_word(request, pk):
    order = get_object_or_404(ServiceOrder, pk=pk)
    doc = Document()
    
    # Márgenes
    s = doc.sections[0]
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(0.4)
    w_total = s.page_width - s.left_margin - s.right_margin

    # Encabezado con Logo
    t = doc.add_table(rows=1, cols=3); t.autofit=False; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width = int(w_total*0.2)
    t.columns[1].width = int(w_total*0.5)
    t.columns[2].width = int(w_total*0.3)
    
    logo = os.path.join(settings.BASE_DIR, 'static', 'orders', 'inovatech-logo.png')
    if os.path.exists(logo):
        t.rows[0].cells[0].paragraphs[0].add_run().add_picture(logo, width=Inches(1.2))
    
    p = t.rows[0].cells[1].paragraphs[0]; p.alignment=1
    r = p.add_run("ORDEN DE SERVICIO TÉCNICO"); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=RGBColor(31,78,120)
    
    # Cuadro de Folio
    tf = t.rows[0].cells[2].add_table(rows=2, cols=1); tf.alignment=2
    tbl = tf._tbl; borders = OxmlElement('w:tblBorders')
    for b in ['top','left','bottom','right','insideH']:
        el = OxmlElement(f'w:{b}'); el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4'); borders.append(el)
    tbl.tblPr.append(borders)
    
    tf.cell(0,0).paragraphs[0].add_run("FOLIO").bold=True
    r = tf.cell(1,0).paragraphs[0].add_run(str(order.folio)); r.bold=True; r.font.color.rgb=RGBColor(192,0,0); r.font.size=Pt(12)
    doc.add_paragraph()

    # Tabla Principal
    t = doc.add_table(rows=0, cols=4); t.style='Table Grid'
    w1=int(w_total*0.15); w2=int(w_total*0.35)
    for i in range(4): t.columns[i].width = w1 if i%2==0 else w2

    # A. Info General
    r=t.add_row(); r.cells[0].merge(r.cells[3]); make_header_blue(r.cells[0], " A. INFORMACIÓN GENERAL")
    r=t.add_row(); make_label_gray(r.cells[0],"1. Cliente"); set_value_text(r.cells[1], order.cliente_nombre)
    make_label_gray(r.cells[2],"2. Ubicación"); set_value_text(r.cells[3], order.ubicacion)
    
    r=t.add_row(); make_label_gray(r.cells[0],"3. Fecha"); set_value_text(r.cells[1], order.fecha_servicio)
    make_label_gray(r.cells[2],"4. Contacto"); set_value_text(r.cells[3], order.cliente_contacto)

    r=t.add_row(); make_label_gray(r.cells[0],"5. Tipo serv.")
    bd = str(order.tipos_servicio).lower()
    chk = lambda x: "☒" if x in bd else "☐"
    txt = (f"{chk('instal')} Instalación   {chk('config')} Configuración\n"
           f"{chk('mantenim')} Mantenimiento   {chk('garant')} Garantía\n"
           f"{chk('revis')} Falla/Revisión   {chk('capacit')} Capacitación")
    set_value_text(r.cells[1], txt)
    make_label_gray(r.cells[2],"6. Ingeniero"); set_value_text(r.cells[3], order.ingeniero_nombre)

    r=t.add_row(); set_cell_color(r.cells[0],'FFFFFF'); set_cell_color(r.cells[1],'FFFFFF')
    make_label_gray(r.cells[2],"7. ID Ticket"); set_value_text(r.cells[3], order.ticket_id)

    # B. Equipo
    r=t.add_row(); r.cells[0].merge(r.cells[3]); make_header_blue(r.cells[0], " B. DATOS DEL EQUIPO")
    r=t.add_row()
    for i,x in enumerate(["Marca","Modelo","Serie","Descripción"]): make_label_gray(r.cells[i], x)
    eq = order.equipos.first()
    r=t.add_row()
    if eq:
        set_value_text(r.cells[0], eq.marca); set_value_text(r.cells[1], eq.modelo)
        set_value_text(r.cells[2], eq.serie); set_value_text(r.cells[3], eq.descripcion)
    else: r.cells[0].merge(r.cells[3]); set_value_text(r.cells[0], "Sin equipo.")

    # C. Datos Técnicos
    r=t.add_row(); r.cells[0].merge(r.cells[3]); make_header_blue(r.cells[0], " C. DATOS TÉCNICOS")
    r=t.add_row(); make_label_gray(r.cells[0],"1. Título"); r.cells[1].merge(r.cells[3]); set_value_text(r.cells[1], order.titulo)
    
    r=t.add_row(); make_label_gray(r.cells[0],"2. Actividades"); c=r.cells[1]; c.merge(r.cells[3]); c.text=""
    c.paragraphs[0].add_run(str(order.actividades)).font.size=Pt(9)
    
    # Evidencias en Word individual
    if order.evidencias.exists():
        c.paragraphs[0].add_run("\n\n--- EVIDENCIA FOTOGRÁFICA ---\n").bold=True
        for f in order.evidencias.all():
            if f.archivo and os.path.exists(f.archivo.path):
                try: c.paragraphs[0].add_run().add_picture(f.archivo.path, width=Inches(2.5))
                except: pass

    r=t.add_row(); make_label_gray(r.cells[0],"3. Comentarios"); r.cells[1].merge(r.cells[3]); set_value_text(r.cells[1], order.comentarios)

    # D. Costos
    r=t.add_row(); r.cells[0].merge(r.cells[3]); make_header_blue(r.cells[0], " D. COSTOS Y TIEMPOS")
    r=t.add_row(); make_label_gray(r.cells[0],"Tiempo (hrs)"); set_value_text(r.cells[1], order.horas)
    make_label_gray(r.cells[2],"Costo"); set_value_text(r.cells[3], str(order.costo_mxn))

    # E. Firmas
    r=t.add_row(); r.cells[0].merge(r.cells[3]); make_header_blue(r.cells[0], " E. ACEPTACIÓN DEL SERVICIO")
    r=t.add_row(); r.cells[0].merge(r.cells[3])
    r.cells[0].paragraphs[0].add_run("Al firmar acepta conformidad.").font.size=Pt(7)

    r=t.add_row(); c=r.cells[0]; c.merge(r.cells[3]); ts=c.add_table(rows=1,cols=3); ts.autofit=False; ts.alignment=1
    w3 = int(w_total/3)
    for cl in ts.rows[0].cells: cl.width = w3

    # Buscador de firmas
    def find_u(n):
        if not n: return None
        target = str(n).lower().strip()
        for u in User.objects.all():
            fn = (u.get_full_name() or "").lower().strip()
            un = u.username.lower().strip()
            if fn == target or un == target: return u
        return None

    fi = None; ui = find_u(order.ingeniero_nombre)
    if ui:
        if hasattr(ui,'engineerprofile') and ui.engineerprofile.firma: fi=ui.engineerprofile.firma
        elif hasattr(ui,'profile') and ui.profile.firma: fi=ui.profile.firma
    
    fv = None; uv = find_u(order.contacto_nombre)
    if uv:
        if hasattr(uv,'engineerprofile') and uv.engineerprofile.firma: fv=uv.engineerprofile.firma
        elif hasattr(uv,'profile') and uv.profile.firma: fv=uv.profile.firma
    
# Firma del Visor (Ejecutivo de Ventas)
    fv = None
    nombre_visor = ""
    if order.visor:
     nombre_visor = order.visor.get_full_name() or order.visor.username
    if hasattr(order.visor, 'profile') and order.visor.profile.firma:
        fv = order.visor.profile.firma

# -----------------------------------------------------------------------------------
    # ESTE BLOQUE DEBE TENER 4 ESPACIOS AL INICIO (ESTAR DENTRO DE LA FUNCIÓN)
    # -----------------------------------------------------------------------------------
    
    insert_signature(ts.cell(0,0), "Ingeniero de Soporte", fi, order.ingeniero_nombre)
    insert_signature(ts.cell(0,1), "Cliente", order.firma, order.cliente_contacto)
    
    # Esta es la línea nueva para el Visor
    insert_signature(ts.cell(0,2), "Contacto Ventas (Visor)", fv, nombre_visor)

    resp = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    resp['Content-Disposition'] = f'attachment; filename="Orden_{order.folio}.docx"'
    doc.save(resp)
    
    return resp